"""
CRT AI Extractor — univerzális horgonypont
Két csatorna:
  PULL  – webshop HTML-ből kinyeri az árakat (llama3:8b)
  PUSH  – szállítói Excel / PDF / Word árlistából kinyeri (llama3:8b)
CSS selector, oldalstruktúra és formátum ismerete nélkül.
"""
import re, json, logging, urllib.request, urllib.parse, os
from scrapers.base import save_price, get_or_create_source, update_source_status

log = logging.getLogger("CRT.scraper.ai")

OLLAMA_URL = "http://localhost:11434/api/generate"
MODEL      = "llama3:8b"

HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                  "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124 Safari/537.36",
    "Accept-Language": "hu-HU,hu;q=0.9",
}


def _fetch(url: str, max_bytes: int = 120_000) -> str:
    import ssl
    ctx = ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE
    req = urllib.request.Request(url, headers=HEADERS)
    with urllib.request.urlopen(req, timeout=15, context=ctx) as r:
        return r.read(max_bytes).decode("utf-8", errors="ignore")


def _html_to_text(html: str) -> str:
    text = re.sub(r'<script[^>]*>.*?</script>', ' ', html, flags=re.S | re.I)
    text = re.sub(r'<style[^>]*>.*?</style>',  ' ', text,  flags=re.S | re.I)
    text = re.sub(r'<[^>]+>', ' ', text)
    text = re.sub(r'&nbsp;', ' ', text)
    text = re.sub(r'&amp;',  '&', text)
    text = re.sub(r'&lt;',   '<', text)
    text = re.sub(r'&gt;',   '>', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def _ask_llama(text_chunk: str) -> str:
    prompt = (
        "Magyar villamos anyag webshop oldal szövege.\n"
        "Listazd az osszes tereket es nettoarakat.\n"
        "Csak valodi termekeket irj, formatum soronkent: TERMEK | AR Ft\n"
        "Ha nincs termek, irj: NINCS\n\n"
        f"{text_chunk}"
    )
    payload = json.dumps({
        "model":   MODEL,
        "prompt":  prompt,
        "stream":  False,
        "options": {"temperature": 0.0, "num_predict": 600},
    }).encode()
    req = urllib.request.Request(
        OLLAMA_URL, data=payload,
        headers={"Content-Type": "application/json"}
    )
    with urllib.request.urlopen(req, timeout=90) as r:
        return json.loads(r.read())["response"]


def _parse_ai_response(text: str) -> list[dict]:
    results = []
    for line in text.splitlines():
        if "|" not in line:
            continue
        parts = line.split("|")
        if len(parts) < 2:
            continue
        name = re.sub(r'^\W+', '', parts[0]).strip()
        price_str = parts[1].strip()
        price_m = re.search(r'(\d[\d\s,.]*)', price_str)
        if not name or not price_m:
            continue
        raw = price_m.group(1).replace('\xa0','').replace(' ','').replace('.','').replace(',','.')
        try:
            price = float(raw)
            if price > 0:
                results.append({"name": name[:200], "price": price})
        except ValueError:
            pass
    return results


def extract_prices(url: str, search_term: str = "") -> list[dict]:
    """Egyetlen URL-ről kinyeri a termékeket AI-val."""
    search_url = url
    if search_term:
        sep = "&" if "?" in url else "?"
        search_url = f"{url}{sep}s={urllib.parse.quote(search_term, safe='+')}"

    log.info("AI fetch: %s", search_url)
    html = _fetch(search_url)
    text = _html_to_text(html)

    # Legfontosabb 5000 karakter megkeresése (ahol az árak vannak)
    best_start = 0
    best_score = 0
    chunk_size = 2000
    for i in range(0, min(len(text), 40000), 1000):
        chunk = text[i: i + chunk_size]
        score = len(re.findall(r'\d+\s*Ft', chunk, re.I))
        if score > best_score:
            best_score = score
            best_start = i

    chunk = text[best_start: best_start + chunk_size]
    log.info("  Legjobb chunk: pos=%d, Ft-ok: %d", best_start, best_score)

    raw_response = _ask_llama(chunk)
    log.debug("AI valasz:\n%s", raw_response)

    products = _parse_ai_response(raw_response)
    log.info("  AI talalt: %d termek", len(products))
    return products


def scrape_source(source_name: str, base_url: str,
                  search_terms: list[str]) -> dict:
    """Teljes scrape egy forrásra, több keresési kifejezéssel."""
    source_id = get_or_create_source(source_name, base_url, "public")
    saved, errors = 0, []

    for term in search_terms:
        try:
            products = extract_prices(base_url, term)
            for p in products:
                save_price(
                    source_id = source_id,
                    raw_name  = f"[AI/{source_name}] {p['name']}",
                    raw_price = p["price"],
                )
                saved += 1
            log.info("%s / '%s': %d mentes", source_name, term, len(products))
        except Exception as e:
            msg = f"{source_name} / '{term}': {e}"
            log.warning(msg)
            errors.append(msg)

    update_source_status(source_id, saved > 0, errors[0] if errors and not saved else None)
    return {"source": source_name, "saved": saved, "errors": errors[:5]}


def _ask_llama_doc(text_chunk: str) -> str:
    """Prompt dokumentumokhoz (Excel/PDF/Word) — rövidebb, tisztább szövegre optimalizálva."""
    prompt = (
        "Magyar szallitoi arlistak adatai.\n"
        "Listazd az osszes termelket es nettoarekat.\n"
        "Csak valodi termekeket irj, formatum soronkent: TERMEKNEV | AR Ft\n"
        "Ha az ar netto, jelold: TERMEKNEV | AR Ft netto\n"
        "Ha nincs ar, hagyd ki a termelket.\n"
        "Ha nincs termek, irj: NINCS\n\n"
        f"{text_chunk}"
    )
    payload = json.dumps({
        "model":   MODEL,
        "prompt":  prompt,
        "stream":  False,
        "options": {"temperature": 0.0, "num_predict": 800},
    }).encode()
    req = urllib.request.Request(
        OLLAMA_URL, data=payload,
        headers={"Content-Type": "application/json"}
    )
    with urllib.request.urlopen(req, timeout=120) as r:
        return json.loads(r.read())["response"]


def extract_from_excel(path: str) -> list[dict]:
    """PUSH csatorna: Excel árlista → termék+ár lista."""
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    lines = []
    for sheet in wb.worksheets:
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                lines.append("  ".join(cells))
        if len(lines) > 2000:
            break
    wb.close()

    text = "\n".join(lines[:2000])
    chunk = text[:4000]
    log.info("Excel feldolgozás: %s (%d sor)", os.path.basename(path), len(lines))
    raw = _ask_llama_doc(chunk)
    log.debug("AI válasz (Excel):\n%s", raw)
    return _parse_ai_response(raw)


def extract_from_pdf(path: str) -> list[dict]:
    """PUSH csatorna: PDF árlista → termék+ár lista."""
    import pdfplumber
    lines = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages[:20]:
            text = page.extract_text() or ""
            lines.extend(text.splitlines())
            if len(lines) > 2000:
                break

    text = "\n".join(lines[:2000])
    chunk = text[:4000]
    log.info("PDF feldolgozás: %s (%d sor)", os.path.basename(path), len(lines))
    raw = _ask_llama_doc(chunk)
    log.debug("AI válasz (PDF):\n%s", raw)
    return _parse_ai_response(raw)


def extract_from_word(path: str) -> list[dict]:
    """PUSH csatorna: Word árlista → termék+ár lista."""
    import docx
    doc = docx.Document(path)
    lines = []

    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append("  |  ".join(cells))

    text = "\n".join(lines[:2000])
    chunk = text[:4000]
    log.info("Word feldolgozás: %s (%d sor)", os.path.basename(path), len(lines))
    raw = _ask_llama_doc(chunk)
    log.debug("AI válasz (Word):\n%s", raw)
    return _parse_ai_response(raw)


def extract_from_file(path: str) -> list[dict]:
    """Automatikus formátum felismerés és kinyerés."""
    ext = os.path.splitext(path)[1].lower()
    if ext in (".xlsx", ".xls", ".xlsm"):
        return extract_from_excel(path)
    elif ext == ".pdf":
        return extract_from_pdf(path)
    elif ext in (".docx", ".doc"):
        return extract_from_word(path)
    else:
        raise ValueError(f"Nem támogatott formátum: {ext}")


if __name__ == "__main__":
    import sys
    logging.basicConfig(level=logging.INFO,
                        format="%(asctime)s | %(levelname)s | %(message)s")

    # ── ÉLES TESZT: govill.hu ─────────────────────────────────
    print("\n=== AI HORGONYPONT — ELES TESZT: govill.hu ===\n")

    # 1. Regex-alapú referencia (mit talál a meglévő scraper)
    from scrapers.govill_hu import _fetch_html, _parse_products
    ref_html  = _fetch_html("https://www.govill.hu/hu/?s=NYM")
    ref_prods = _parse_products(ref_html)
    print(f"Regex referencia: {len(ref_prods)} termek")
    for p in ref_prods[:5]:
        print(f"  {p['name'][:50]:50s} | {p['price']:>10,.0f} Ft")

    # 2. AI extrakció ugyanarról az URL-ről
    print(f"\nAI extrakció (llama3:8b)...")
    ai_prods = extract_prices("https://www.govill.hu/hu/", "NYM")
    print(f"AI talalt: {len(ai_prods)} termek")
    for p in ai_prods[:5]:
        print(f"  {p['name'][:50]:50s} | {p['price']:>10,.0f} Ft")

    # 3. Pontossági mérés
    print("\n--- Pontossag ---")
    matched = 0
    for ai_p in ai_prods:
        for ref_p in ref_prods:
            price_ok = abs(ai_p["price"] - ref_p["price"]) < 10
            name_ok  = any(w in ai_p["name"].lower()
                           for w in ref_p["name"].lower().split()[:3] if len(w) > 3)
            if price_ok or name_ok:
                matched += 1
                break

    total = max(len(ref_prods), 1)
    pct   = matched / total * 100
    print(f"Egyezes: {matched}/{total} = {pct:.0f}%")
    print(f"Eredmeny: {'MEGFELEL (>=90%)' if pct >= 90 else 'NEM ERI EL'}")
    if pct >= 90:
        print("\n→ ai_extractor.py PRODUKTIONBA KERUL")
