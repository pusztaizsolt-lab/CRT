"""
CRT Ajánlatsegéd – Export router v0.8
Végpontok: GET /quotes/{id}/export?format=xlsx|docx|pdf
Függőségek: openpyxl · python-docx · reportlab
  py -3.11 -m pip install openpyxl python-docx reportlab
"""
from fastapi import APIRouter, HTTPException, Depends, Query
from fastapi.responses import StreamingResponse
from sqlalchemy import create_engine, text
from auth import require_auth
import io, os, logging
from env_detect import get_db_url
from datetime import datetime, timezone, timedelta

router = APIRouter(tags=["export"])
log    = logging.getLogger("CRT.export")

DB_URL = get_db_url()
engine = create_engine(DB_URL, pool_pre_ping=True, pool_size=2)

CET  = timezone(timedelta(hours=1))
CEST = timezone(timedelta(hours=2))

def _local_now():
    utc = datetime.utcnow()
    tz  = CEST if 3 <= utc.month <= 10 else CET
    return utc.replace(tzinfo=timezone.utc).astimezone(tz)

def _get_config(key: str) -> str:
    try:
        with engine.connect() as conn:
            row = conn.execute(text("SELECT value FROM system_config WHERE key=:k"), {"k": key}).fetchone()
            return row[0] if row else ""
    except Exception:
        return ""

def _load_quote(quote_id: int) -> dict:
    with engine.connect() as conn:
        q = conn.execute(text(
            "SELECT id, quote_number, title, client_name, client_ref, status, "
            "valid_until, notes, created_at FROM quotes WHERE id=:id"
        ), {"id": quote_id}).fetchone()
        if not q:
            raise HTTPException(404, "Ajánlat nem található")

        lines = conn.execute(text(
            "SELECT line_no, name, manufacturer, unit, quantity, unit_price, "
            "total_price, cell_status, item_type "
            "FROM quote_lines WHERE quote_id=:id ORDER BY line_no"
        ), {"id": quote_id}).fetchall()

    q_dict = dict(q._mapping)
    for k, v in q_dict.items():
        if hasattr(v, "isoformat"):
            q_dict[k] = v.isoformat()

    l_list = []
    for l in lines:
        d = dict(l._mapping)
        for k, v in d.items():
            if hasattr(v, "isoformat"):
                d[k] = v.isoformat()
        l_list.append(d)

    return q_dict, l_list


# ── EXCEL EXPORT ──────────────────────────────────────────────

def _build_xlsx(quote: dict, lines: list) -> bytes:
    try:
        import openpyxl
        from openpyxl.styles import (
            Font, PatternFill, Alignment, Border, Side, numbers
        )
        from openpyxl.utils import get_column_letter
    except ImportError:
        raise HTTPException(500, "openpyxl hiányzik – telepítés: py -3.11 -m pip install openpyxl")

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Ajánlat"

    company = _get_config("company_name") or "Civil Rendszertechnika Kft."
    company_tax = _get_config("company_tax") or ""

    # ── Stílusok ──
    hdr_fill   = PatternFill("solid", fgColor="1A1A2E")
    sub_fill   = PatternFill("solid", fgColor="16213E")
    row_fill   = PatternFill("solid", fgColor="0F3460")
    alt_fill   = PatternFill("solid", fgColor="0A2744")
    tot_fill   = PatternFill("solid", fgColor="1A1A2E")
    white_font = Font(color="E8E8F0", bold=True)
    dim_font   = Font(color="9898B8")
    acc_font   = Font(color="7C6AF7", bold=True)
    thin = Side(style="thin", color="2A2A3E")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    center = Alignment(horizontal="center", vertical="center")
    right  = Alignment(horizontal="right",  vertical="center")
    wrap   = Alignment(wrap_text=True, vertical="center")

    # ── Fejléc blokk ──
    ws.merge_cells("A1:H1")
    ws["A1"] = company
    ws["A1"].font = Font(color="7C6AF7", bold=True, size=14)
    ws["A1"].fill = hdr_fill

    ws.merge_cells("A2:H2")
    ws["A2"] = f"Ajánlatszám: {quote.get('quote_number','–')}  |  Dátum: {_local_now().strftime('%Y.%m.%d')}"
    ws["A2"].font = dim_font
    ws["A2"].fill = hdr_fill

    ws.merge_cells("A3:D3"); ws["A3"] = f"Ügyfél: {quote.get('client_name','–')}"; ws["A3"].fill = sub_fill; ws["A3"].font = white_font
    ws.merge_cells("E3:H3"); ws["E3"] = f"Hivatkozás: {quote.get('client_ref','–')}"; ws["E3"].fill = sub_fill; ws["E3"].font = dim_font

    ws.merge_cells("A4:H4"); ws["A4"] = f"Tárgy: {quote.get('title','–')}"; ws["A4"].fill = sub_fill; ws["A4"].font = white_font

    # ── Táblázat fejléc ──
    headers = ["#", "Megnevezés", "Gyártó", "Me.", "Menny.", "Egységár (Ft)", "Összesen (Ft)", "Státusz"]
    widths  = [5,    45,           18,        6,     10,        16,              16,               12]
    for i, (h, w) in enumerate(zip(headers, widths), 1):
        col = get_column_letter(i)
        ws.column_dimensions[col].width = w
        cell = ws[f"{col}6"]
        cell.value = h
        cell.font  = white_font
        cell.fill  = row_fill
        cell.alignment = center if i in (1,4,5) else (right if i in (6,7) else Alignment(vertical="center"))
        cell.border = border

    ws.row_dimensions[6].height = 22

    # ── Sorok ──
    status_labels = {
        "done": "Kész", "identified": "Azonosított", "uncertain": "Bizonytalan",
        "raw": "Nem azonosított", "manual": "Kézzel", "skip": "Kihagyva"
    }
    grand_total = 0.0

    for row_i, l in enumerate(lines, 7):
        fill = alt_fill if row_i % 2 == 0 else PatternFill("solid", fgColor="122244")
        total = l.get("total_price") or (
            (l.get("unit_price") or 0) * (l.get("quantity") or 0)
        )
        grand_total += float(total or 0)

        values = [
            l.get("line_no", row_i - 6),
            l.get("name") or l.get("raw_name") or "–",
            l.get("manufacturer") or "",
            l.get("unit") or "db",
            l.get("quantity") or 1,
            l.get("unit_price"),
            total or None,
            status_labels.get(l.get("cell_status", "raw"), l.get("cell_status", "–")),
        ]
        aligns = [center, wrap, Alignment(vertical="center"), center, center, right, right, center]

        for ci, (val, aln) in enumerate(zip(values, aligns), 1):
            c = ws.cell(row=row_i, column=ci, value=val)
            c.fill = fill; c.alignment = aln; c.border = border
            c.font = Font(color="D8D8E8", size=10)
            if ci in (6, 7) and val is not None:
                c.number_format = '#,##0'

    # ── Összesítő sor ──
    tot_row = len(lines) + 7
    ws.merge_cells(f"A{tot_row}:F{tot_row}")
    ws[f"A{tot_row}"] = "VÉGÖSSZEG"
    ws[f"A{tot_row}"].font = Font(color="7C6AF7", bold=True)
    ws[f"A{tot_row}"].fill = tot_fill
    ws[f"A{tot_row}"].alignment = right

    ws[f"G{tot_row}"] = grand_total
    ws[f"G{tot_row}"].font = Font(color="3ECF8E", bold=True, size=12)
    ws[f"G{tot_row}"].fill = tot_fill
    ws[f"G{tot_row}"].alignment = right
    ws[f"G{tot_row}"].number_format = '#,##0'
    ws[f"G{tot_row}"].border = border

    # ── Megjegyzés ──
    note_row = tot_row + 2
    validity = _get_config("quote_validity_days") or "30"
    ws.merge_cells(f"A{note_row}:H{note_row}")
    ws[f"A{note_row}"] = f"Az ajánlat {validity} napig érvényes. | {company_tax}"
    ws[f"A{note_row}"].font = Font(color="5A5A7A", italic=True, size=9)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


# ── WORD EXPORT ───────────────────────────────────────────────

def _build_docx(quote: dict, lines: list) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import copy
    except ImportError:
        raise HTTPException(500, "python-docx hiányzik – telepítés: py -3.11 -m pip install python-docx")

    company     = _get_config("company_name") or "Civil Rendszertechnika Kft."
    company_tax = _get_config("company_tax")  or ""
    validity    = _get_config("quote_validity_days") or "30"

    doc = Document()

    # Margók
    for sec in doc.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    def add_run(para, text, bold=False, size=11, color=None):
        run = para.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        return run

    # ── Fejléc ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, company, bold=True, size=16, color=(124,106,247))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, f"ÁRAJÁNLAT – {quote.get('quote_number','–')}", bold=True, size=13)

    doc.add_paragraph()

    # ── Meta blokk ──
    meta = [
        ("Ügyfél",       quote.get("client_name") or "–"),
        ("Hivatkozás",   quote.get("client_ref")  or "–"),
        ("Tárgy",        quote.get("title")        or "–"),
        ("Dátum",        _local_now().strftime("%Y.%m.%d")),
        ("Érvényesség",  f"{validity} nap"),
    ]
    for label, val in meta:
        p = doc.add_paragraph()
        add_run(p, f"{label}: ", bold=True, size=10, color=(152,152,184))
        add_run(p, val, size=10)

    doc.add_paragraph()

    # ── Tábla ──
    headers = ["#", "Megnevezés", "Gyártó", "Me.", "Menny.", "Egységár (Ft)", "Összesen (Ft)"]
    col_w   = [Cm(1), Cm(7.5), Cm(3), Cm(1.2), Cm(1.8), Cm(3), Cm(3)]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Fejlécsor
    hdr_row = table.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_w)):
        cell = hdr_row.cells[i]
        cell.width = w
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(232,232,240)

    grand_total = 0.0
    for l in lines:
        total = l.get("total_price") or (
            (l.get("unit_price") or 0) * (l.get("quantity") or 0)
        )
        grand_total += float(total or 0)

        row = table.add_row()
        vals = [
            str(l.get("line_no","")),
            l.get("name") or l.get("raw_name") or "–",
            l.get("manufacturer") or "",
            l.get("unit") or "db",
            str(l.get("quantity") or 1),
            f"{int(l.get('unit_price') or 0):,}".replace(",","  ") if l.get("unit_price") else "–",
            f"{int(total or 0):,}".replace(",","  ") if total else "–",
        ]
        aligns = [
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.RIGHT,
            WD_ALIGN_PARAGRAPH.RIGHT,
            WD_ALIGN_PARAGRAPH.RIGHT,
        ]
        for ci, (val, aln) in enumerate(zip(vals, aligns)):
            c = row.cells[ci]
            c.paragraphs[0].alignment = aln
            r = c.paragraphs[0].add_run(val)
            r.font.size = Pt(9)

    # Összesítő sor
    tot_row = table.add_row()
    tot_row.cells[0].merge(tot_row.cells[5])
    p = tot_row.cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("VÉGÖSSZEG")
    run.bold = True; run.font.size = Pt(10)

    p2 = tot_row.cells[6].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = p2.add_run(f"{int(grand_total):,}".replace(","," ") + " Ft")
    run2.bold = True; run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(62,207,142)

    doc.add_paragraph()

    # ── Megjegyzés ──
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(note, f"Az ajánlat {validity} napig érvényes. {company_tax}", size=8, color=(90,90,122))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── PDF EXPORT ────────────────────────────────────────────────

def _build_pdf(quote: dict, lines: list) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Table, TableStyle, Paragraph,
            Spacer, HRFlowable,
        )
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        raise HTTPException(500, "reportlab hiányzik – telepítés: py -3.11 -m pip install reportlab")

    company  = _get_config("company_name") or "Civil Rendszertechnika Kft."
    validity = _get_config("quote_validity_days") or "30"
    tax      = _get_config("company_tax") or ""

    # Szín paletta (ejszaka téma)
    C_BG    = colors.HexColor("#16161f")
    C_ACC   = colors.HexColor("#7c6af7")
    C_TEXT  = colors.HexColor("#e8e8f0")
    C_DIM   = colors.HexColor("#9898b8")
    C_GREEN = colors.HexColor("#3ecf8e")
    C_ROW1  = colors.HexColor("#0f3460")
    C_ROW2  = colors.HexColor("#0a2744")
    C_HDR   = colors.HexColor("#1a1a2e")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm,
    )

    styles = getSampleStyleSheet()
    s_company = ParagraphStyle("company",  fontSize=16, textColor=C_ACC,   spaceAfter=4,  alignment=1)
    s_title   = ParagraphStyle("title",    fontSize=13, textColor=C_TEXT,  spaceAfter=6,  alignment=1, fontName="Helvetica-Bold")
    s_meta    = ParagraphStyle("meta",     fontSize=9,  textColor=C_DIM,   spaceAfter=2)
    s_footer  = ParagraphStyle("footer",   fontSize=8,  textColor=C_DIM,   alignment=1)

    story = []

    # Fejléc
    story.append(Paragraph(company, s_company))
    story.append(Paragraph(f"ÁRAJÁNLAT – {quote.get('quote_number','–')}", s_title))
    story.append(HRFlowable(width="100%", thickness=1, color=C_ACC, spaceAfter=8))

    # Meta adatok
    meta_pairs = [
        ("Ügyfél",     quote.get("client_name") or "–"),
        ("Hivatkozás", quote.get("client_ref")  or "–"),
        ("Tárgy",      quote.get("title")       or "–"),
        ("Dátum",      _local_now().strftime("%Y.%m.%d")),
        ("Érvényes",   f"{validity} napig"),
    ]
    for label, val in meta_pairs:
        story.append(Paragraph(f"<b>{label}:</b>  {val}", s_meta))
    story.append(Spacer(1, 0.4*cm))

    # Tábla fejléc + sorok
    headers = ["#", "Megnevezés", "Gyártó", "Me.", "Menny.", "Egységár (Ft)", "Összesen (Ft)"]
    col_w   = [1*cm, 6.5*cm, 3*cm, 1.2*cm, 1.8*cm, 2.8*cm, 2.8*cm]

    grand_total = 0.0
    table_data  = [headers]

    for l in lines:
        total = l.get("total_price") or (
            (l.get("unit_price") or 0) * (l.get("quantity") or 0)
        )
        grand_total += float(total or 0)
        table_data.append([
            str(l.get("line_no", "")),
            l.get("name") or l.get("raw_name") or "–",
            l.get("manufacturer") or "",
            l.get("unit") or "db",
            str(l.get("quantity") or 1),
            f"{int(l.get('unit_price') or 0):,}".replace(",", " ") if l.get("unit_price") else "–",
            f"{int(total or 0):,}".replace(",", " ") if total else "–",
        ])

    # Összesítő sor
    table_data.append(["", "", "", "", "", "VÉGÖSSZEG",
                        f"{int(grand_total):,}".replace(",", " ") + " Ft"])

    t = Table(table_data, colWidths=col_w, repeatRows=1)
    last = len(table_data) - 1
    t.setStyle(TableStyle([
        # fejléc
        ("BACKGROUND",  (0, 0), (-1,  0), C_HDR),
        ("TEXTCOLOR",   (0, 0), (-1,  0), C_TEXT),
        ("FONTNAME",    (0, 0), (-1,  0), "Helvetica-Bold"),
        ("FONTSIZE",    (0, 0), (-1,  0), 8),
        ("ALIGN",       (0, 0), (-1,  0), "CENTER"),
        # sorok (zebra)
        ("BACKGROUND",  (0, 1), (-1, last-1), C_ROW1),
        ("ROWBACKGROUNDS", (0, 1), (-1, last-1), [C_ROW1, C_ROW2]),
        ("TEXTCOLOR",   (0, 1), (-1, last-1), C_TEXT),
        ("FONTSIZE",    (0, 1), (-1, last-1), 8),
        ("ALIGN",       (4, 1), (-1, last-1), "RIGHT"),
        # összesítő sor
        ("BACKGROUND",  (0, last), (-1, last), C_HDR),
        ("TEXTCOLOR",   (5, last), (5, last), C_DIM),
        ("TEXTCOLOR",   (6, last), (6, last), C_GREEN),
        ("FONTNAME",    (5, last), (-1, last), "Helvetica-Bold"),
        ("FONTSIZE",    (5, last), (-1, last), 9),
        ("ALIGN",       (5, last), (-1, last), "RIGHT"),
        # keret
        ("GRID",        (0, 0), (-1, -1), 0.3, colors.HexColor("#2a2a3e")),
        ("TOPPADDING",  (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(t)

    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph(f"Az ajánlat {validity} napig érvényes.  {tax}", s_footer))

    doc.build(story)
    buf.seek(0)
    return buf.read()


# ── ENDPOINT ──────────────────────────────────────────────────

@router.get("/quotes/{quote_id}/export")
async def export_quote(
    quote_id: int,
    format:   str = Query("xlsx", regex="^(xlsx|docx|pdf)$"),
    _auth:    dict = Depends(require_auth),
):
    """Ajánlat exportálása Excel vagy Word formátumban"""
    quote, lines = _load_quote(quote_id)

    qno = (quote.get("quote_number") or f"ajanalt_{quote_id}").replace("/", "-")

    if format == "xlsx":
        data     = _build_xlsx(quote, lines)
        filename = f"{qno}.xlsx"
        media    = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    elif format == "pdf":
        data     = _build_pdf(quote, lines)
        filename = f"{qno}.pdf"
        media    = "application/pdf"
    else:
        data     = _build_docx(quote, lines)
        filename = f"{qno}.docx"
        media    = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    log.info("Export: quote_id=%s format=%s user=%s", quote_id, format, _auth.get("username"))

    try:
        with engine.begin() as conn:
            conn.execute(text(
                "INSERT INTO audit_log (log_id, user_id, action, entity_id, description, timestamp) "
                "VALUES (:lid, :uid, 'export', :eid, :desc, NOW())"
            ), {
                "lid":  str(__import__('uuid').uuid4()),
                "uid":  str(_auth.get("user_id")),
                "eid":  str(quote_id),
                "desc": f"Export: {filename}",
            })
    except Exception:
        pass

    return StreamingResponse(
        io.BytesIO(data),
        media_type=media,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
