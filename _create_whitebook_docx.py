#!/usr/bin/env python3
"""
CRT Whitebook v1.0 – Word dokumentum generátor
Futtatás: py -3.11 _create_whitebook_docx.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUT = "CRT_Whitebook_v1.0.docx"

# ── SZÍN PALETTA ──────────────────────────────────────────────
C_DARK   = RGBColor(0x0d, 0x0d, 0x14)   # háttér
C_ACCENT = RGBColor(0x7c, 0x6a, 0xf7)   # lila
C_DONE   = RGBColor(0x3e, 0xcf, 0x8e)   # zöld
C_WARN   = RGBColor(0xf5, 0xa6, 0x23)   # narancs
C_TEXT   = RGBColor(0xe8, 0xe8, 0xf0)   # világos
C_DIM    = RGBColor(0x60, 0x60, 0x80)   # szürke
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_HEAD_BG = RGBColor(0x16, 0x16, 0x1f)  # sötét kártya
C_ROW_BG  = RGBColor(0x1a, 0x1a, 0x2e)
C_ALT_BG  = RGBColor(0x12, 0x12, 0x20)


def rgb_hex(rgb) -> str:
    return '%02X%02X%02X' % (rgb[0], rgb[1], rgb[2])


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(rgb))
    tcPr.append(shd)


def set_cell_border(cell, side='bottom', color='2a2a3e', sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    border = OxmlElement(f'w:{side}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), str(sz))
    border.set(qn('w:space'), '0')
    border.set(qn('w:color'), color)
    tcBorders.append(border)


def no_space_before(para):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '60')
    pPr.append(spacing)


def add_heading(doc, text, level=1):
    if level == 1:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr = para._p.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '280')
        spacing.set(qn('w:after'), '100')
        pPr.append(spacing)
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = C_ACCENT
        run.font.name = 'Calibri'
        # Vonal alá
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), rgb_hex(C_ACCENT))
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        para = doc.add_paragraph()
        pPr = para._p.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '180')
        spacing.set(qn('w:after'), '60')
        pPr.append(spacing)
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = C_DONE
        run.font.name = 'Calibri'
    elif level == 3:
        para = doc.add_paragraph()
        pPr = para._p.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '120')
        spacing.set(qn('w:after'), '40')
        pPr.append(spacing)
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = C_WARN
        run.font.name = 'Calibri'
    return para


def add_body(doc, text, italic=False, color=None):
    para = doc.add_paragraph()
    no_space_before(para)
    run = para.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.italic = italic
    run.font.color.rgb = color or C_DIM
    return para


def add_code(doc, text):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '40')
    spacing.set(qn('w:after'), '40')
    pPr.append(spacing)
    # Szürke háttér
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1a1a2e')
    pPr.append(shd)
    # Bal margó
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '400')
    pPr.append(ind)
    for line in text.split('\n'):
        run = para.add_run(line + '\n')
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = C_DONE
    return para


def add_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'

    # Fejléc sor
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, C_HEAD_BG)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = C_ACCENT
        run.font.name = 'Calibri'
        for side in ('top', 'bottom', 'left', 'right'):
            set_cell_border(cell, side, color='3d3478', sz=4)

    # Adat sorok
    for ri, row in enumerate(rows):
        bg = C_ROW_BG if ri % 2 == 0 else C_ALT_BG
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            set_cell_bg(cell, bg)
            para = cell.paragraphs[0]
            run = para.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            run.font.color.rgb = C_TEXT
            for side in ('top', 'bottom', 'left', 'right'):
                set_cell_border(cell, side, color='2a2a3e', sz=2)

    # Oszlopszélesség
    if col_widths:
        for row in table.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)

    doc.add_paragraph()  # térköz
    return table


# ══════════════════════════════════════════════════════════════
#  DOKUMENTUM
# ══════════════════════════════════════════════════════════════

doc = Document()

# Oldalbeállítás
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)

# Háttérszín az egész dokumentumra (body background)
body = doc.element.body
sectPr = body.find(qn('w:sectPr'))
if sectPr is None:
    sectPr = OxmlElement('w:sectPr')
    body.append(sectPr)


# ── FEDŐLAP ────────────────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = title_para._p.get_or_add_pPr()
spacing = OxmlElement('w:spacing')
spacing.set(qn('w:before'), '600')
spacing.set(qn('w:after'), '120')
pPr.append(spacing)
r = title_para.add_run('CRT Ajánlatsegéd')
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = C_ACCENT
r.font.name = 'Calibri'

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_para.add_run('Whitebook v1.0')
r2.font.size = Pt(16)
r2.font.color.rgb = C_DONE
r2.font.name = 'Calibri'
r2.bold = True

doc.add_paragraph()
comp_para = doc.add_paragraph()
comp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = comp_para.add_run('Civil Rendszertechnika Kft.\nHelyi AI alapú ajánlatkezelő rendszer')
r3.font.size = Pt(11)
r3.font.color.rgb = C_DIM
r3.font.name = 'Calibri'
r3.italic = True

doc.add_paragraph()
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = date_para.add_run(f'Utolsó frissítés: 2026-06-10  ·  Verzió: 1.0 (teljes)')
r4.font.size = Pt(9)
r4.font.color.rgb = C_DIM
r4.font.name = 'Calibri'

doc.add_page_break()


# ── 1. RENDSZER ÁTTEKINTÉS ─────────────────────────────────────
add_heading(doc, '1. Rendszer áttekintés', 1)
add_body(doc,
    'A CRT Ajánlatsegéd egy helyi, offline-képes AI alapú ajánlatkészítő szoftver, '
    'amelyet a Civil Rendszertechnika Kft. villamossági és épületautomatizálási '
    'ajánlatkészítési folyamatának automatizálására fejlesztettünk.',
    color=C_TEXT)

add_heading(doc, 'Fő értékek', 2)
add_table(doc,
    ['Funkció', 'Leírás'],
    [
        ['AI azonosítás',      'Ajánlatkérő dokumentumból automatikusan azonosítja a tételeket'],
        ['Árfigyelés',         'Webes scraping + kézi ár bevitel, forrás tracking'],
        ['Ajánlatkészítés',    '3 indítási mód, cell-szintű státuszkövetés, xlsx/docx/PDF export'],
        ['Helyi LLM',          'Ollama (llama3:8b / mistral:7b) – internet nélkül is működik'],
        ['LoRA finomhangolás', 'Saját jóváhagyott példákból fine-tuned AI modell'],
        ['Rajz elemzés',       'LLaVA Vision alapú tervrajz/kép elemzés'],
        ['Adatbiztonság',      'Minden adat helyi gépen, titkosított auth, audit log'],
    ],
    col_widths=[5, 12]
)


# ── 2. ARCHITEKTÚRA ────────────────────────────────────────────
add_heading(doc, '2. Architektúra', 1)
add_heading(doc, 'Telepítési modell (ajánlott: WSL2)', 2)
add_code(doc,
'Windows szerver (Win10/11, x64, 250GB SSD)\n'
'│\n'
'├─ start.bat ──→ wsl -d CRT ──→ wsl_start.sh\n'
'│\n'
'└─── WSL2 Ubuntu 22.04 "CRT" disztribúció ──────────────────┐\n'
'     │  PostgreSQL 16      localhost:5432                   │\n'
'     │  FastAPI/uvicorn    localhost:8000  (2 worker)       │\n'
'     │  ChromaDB           localhost:8001                   │\n'
'     │  Ollama LLM         localhost:11434                  │\n'
'     │  Nginx              localhost:80 (UI)                │\n'
'     └────────────────────────────────────────────────────────┘\n'
'          ↑ böngésző: http://localhost'
)

add_heading(doc, 'Helyigény (250 GB SSD)', 2)
add_table(doc,
    ['Komponens', 'Méret'],
    [
        ['WSL2 Ubuntu base', '~1.2 GB'],
        ['Python + csomagok (LoRA-val)', '~1.2 GB'],
        ['Playwright Chromium', '~350 MB'],
        ['llama3:8b + mistral:7b', '~9.5 GB'],
        ['llava:7b (Vision)', '~4.1 GB'],
        ['ChromaDB + vektorok', '~500 MB'],
        ['LoRA adapter', '~500 MB'],
        ['CRT kód + UI', '~50 MB'],
        ['PostgreSQL adatok', '~100 MB'],
        ['ÖSSZESEN', '~17.5 GB (szabad: ~232 GB)'],
    ],
    col_widths=[8, 4]
)


# ── 3. TECHNOLÓGIAI STACK ──────────────────────────────────────
add_heading(doc, '3. Technológiai stack', 1)
add_table(doc,
    ['Réteg', 'Technológia'],
    [
        ['Backend',          'Python 3.11 · FastAPI 0.115 · SQLAlchemy 2.0'],
        ['Adatbázis',        'PostgreSQL 16'],
        ['Auth',             'bcrypt PIN + email OTP + JWT HS256 (8h)'],
        ['Biztonság',        'SanitizeMiddleware (SQL inj / XSS / path traversal)'],
        ['AI – fő',          'Claude API (claude-sonnet-4-6)'],
        ['AI – LoRA',        'HuggingFace transformers + PEFT + TRL (Phi-3-mini / TinyLlama)'],
        ['AI – helyi',       'Ollama (llama3:8b, mistral:7b, llava:7b)'],
        ['Vektoros keresés', 'ChromaDB + sentence-transformers'],
        ['Web scraping',     'Playwright (Chromium)'],
        ['Frontend',         'Vanilla HTML/CSS/JS – 4 CSS téma, nincs Node/npm'],
        ['Export',           'openpyxl (xlsx) · python-docx (docx) · reportlab (PDF)'],
        ['OCR',              'pdfplumber → pytesseract + Pillow (fallback)'],
    ],
    col_widths=[5, 12]
)


# ── 4. ADATBÁZIS SÉMA ─────────────────────────────────────────
add_heading(doc, '4. Adatbázis séma', 1)
add_table(doc,
    ['Tábla', 'Csoport', 'Leírás'],
    [
        ['products',        'Cikktörzs', 'Termékek (item_id UUID, crt_code, manufacturer, unit, status)'],
        ['activities',      'Cikktörzs', 'Tevékenységek (activity_id UUID, unit_type)'],
        ['categories',      'Cikktörzs', 'Kategória fa (parent_id, item_class, sort_order)'],
        ['prices',          'Árak',      'Ármozgások (price_type, source_id, supplier_code, currency)'],
        ['quotes',          'Ajánlatok', 'Fejlécek (quote_number: CRT-YYYY-NNNN, source_mode, client_ref)'],
        ['quote_lines',     'Ajánlatok', 'Tételsorok (line_no, raw_name, cell_status, confidence)'],
        ['web_sources',     'Scraping',  'Forrás URL-ek (base_url, source_type, status)'],
        ['web_scripts',     'Scraping',  'Playwright scriptek (steps JSONB, version, active)'],
        ['web_prices',      'Scraping',  'Scrapelt árak (source_id, item_id, price, currency)'],
        ['golden_examples', 'AI',        'Jóváhagyott azonosítások (raw_text, clean_name, source)'],
        ['chroma_index',    'AI',        'ChromaDB vektor indexek (collection, vector_id)'],
        ['lora_jobs',       'AI',        'LoRA tréning futtatások (status, train_loss, adapter_path)'],
        ['users',           'Rendszer',  'Felhasználók (pin_hash, email, role, locked_until)'],
        ['system_config',   'Rendszer',  'Kulcs-érték konfiguráció'],
        ['audit_log',       'Rendszer',  'Minden esemény naplója (action, entity_type, ip_address)'],
    ],
    col_widths=[4, 3, 10]
)

add_heading(doc, 'Migrációs sorrend (kötelező, friss gépen)', 2)
add_code(doc,
'py -3.11 db_migrate_v02.py   # cikktörzs + árak + audit + config\n'
'py -3.11 db_migrate_v04.py   # auth: users + auth_tokens\n'
'py -3.11 db_migrate_v05.py   # web_sources, web_scripts, web_prices, quotes\n'
'py -3.11 db_migrate_v06.py   # IF NOT EXISTS bővítések + indexek\n'
'py -3.11 db_migrate_v07.py   # golden_examples + chroma_index + Ollama config\n'
'py -3.11 db_migrate_v08.py   # lora_jobs + LoRA system_config kulcsok'
)


# ── 5. AI PIPELINE ────────────────────────────────────────────
add_heading(doc, '5. AI pipeline', 1)
add_heading(doc, 'Azonosítási prioritási lánc', 2)
add_code(doc,
'POST /cikktorzs/identify\n'
'        │\n'
'        ▼\n'
'   1. Claude API ────────── ha claude_api_key van a system_config-ban\n'
'        │ hiba / nincs kulcs\n'
'        ▼\n'
'   2. LoRA fine-tuned ───── ha lora_active_job_id + lora_adapter_path be van állítva\n'
'        │ hiba / nincs aktiválva\n'
'        ▼\n'
'   3. Ollama helyi LLM ──── http://localhost:11434 (llama3:8b alapértelmezett)\n'
'        │ hiba\n'
'        ▼\n'
'   RuntimeError → 503 válasz'
)

add_heading(doc, 'LoRA finomhangolás folyamata', 2)
add_table(doc,
    ['Lépés', 'Leírás'],
    [
        ['1. Adatgyűjtés',   'acceptAll() az ajánlatkészítőben → golden_examples tábla'],
        ['2. Tréning indítás', 'ui/lora.html → minimum 10 jóváhagyott példa szükséges'],
        ['3. GPU mód',       'Phi-3-mini-4k-instruct + 4-bit QLoRA (~30 perc)'],
        ['4. CPU mód',       'TinyLlama-1.1B + fp32 (~2–4 óra 50 példán)'],
        ['5. Adapter mentés', 'models/lora/{job_id}/adapter/'],
        ['6. Aktiválás',     'lora.html → ai_motor.py betölti és cachelve tartja'],
    ],
    col_widths=[5, 12]
)


# ── 6. FRONTEND OLDALAK ───────────────────────────────────────
add_heading(doc, '6. Frontend oldalak (13 db)', 1)
add_table(doc,
    ['Oldal', 'Fájl', 'Leírás'],
    [
        ['Bejelentkezés',    'login.html',          'PIN → OTP 2 lépés, 4 téma, brute force visszaszámláló'],
        ['Főmenü',           'fomenu.html',         '10 funkciókártya, státusz chip-ek'],
        ['Admin konzol',     'admin.html',          'Felhasználókezelés, lock/unlock, PIN reset'],
        ['Cikktörzs',        'cikktorzs.html',      'Fa struktúra, keresés, fájl feltöltés, AI azonosítás'],
        ['Ajánlatkészítő',   'ajanlatkezelo.html',  '3 mód, cell-státusz, AI, xlsx/docx/PDF export'],
        ['Webes árak',       'webes_arak.html',     'Árforrások, script rögzítés, Tanít gomb'],
        ['Beállítások',      'beallitasok.html',    'SMTP, Claude/Ollama kulcsok, PIN csere, backup'],
        ['Naplók',           'naplok.html',         'Audit log, szűrők, CSV export'],
        ['Árak',             'arak.html',           'Ár lista, kézi bevitel, stat sáv'],
        ['Widget',           'widget.html',         'Élő rendszerállapot (API/DB/Ollama/ChromaDB)'],
        ['Rajzelemző',       'rajz_elemzo.html',    'LLaVA Vision, drag&drop kép/PDF, eredmény kártyák'],
        ['LoRA',             'lora.html',           'Tréning indítás, live progress, aktiválás (v1.0)'],
        ['Diagnosztika',     'kezelőpult_v2.html',  'Rendszer diagnosztika, palettaszerkesztő'],
    ],
    col_widths=[3.5, 4, 9.5]
)


# ── 7. BIZTONSÁG ──────────────────────────────────────────────
add_heading(doc, '7. Biztonság', 1)
add_table(doc,
    ['Mechanizmus', 'Leírás'],
    [
        ['PIN',             'bcrypt hash, 6 számjegy'],
        ['2FA (OTP)',       'Email OTP, 10 perc érvényes; SMTP nélkül naplóba kerül (dev mód)'],
        ['JWT',             'HS256, 8 óra, CRT_JWT_SECRET env var'],
        ['Brute force',     '4 kísérlet → 30 perces automatikus zárolás'],
        ['SQL injection',   'SanitizeMiddleware – UNION, DROP, SELECT* stb. blokkolva'],
        ['XSS',             '<script>, javascript:, event handler blokkolva'],
        ['Path traversal',  '../ és ..\\ blokkolva'],
        ['Méret limit',     '512 KB kérés méret korlát'],
        ['Audit log',       'Minden esemény rögzítve: user, IP, entitás, leírás'],
    ],
    col_widths=[5, 12]
)


# ── 8. TELEPÍTÉSI ÚTMUTATÓ ─────────────────────────────────────
add_heading(doc, '8. Telepítési útmutató', 1)
add_heading(doc, 'Gyors indítás (fejlesztési mód, Windows natív)', 2)
add_code(doc,
'# 1. Csomagok\n'
'py -3.11 -m pip install -r requirements.txt\n\n'
'# 2. DB inicializálás (sorban!)\n'
'py -3.11 db_migrate_v02.py && py -3.11 db_migrate_v04.py\n'
'py -3.11 db_migrate_v05.py && py -3.11 db_migrate_v06.py\n'
'py -3.11 db_migrate_v07.py && py -3.11 db_migrate_v08.py\n\n'
'# 3. Admin felhasználó\n'
'py -3.11 _setup/create_admin.py\n\n'
'# 4. Backend\n'
'py -3.11 -m uvicorn main:app --reload\n\n'
'# 5. Böngésző: ui/login.html'
)

add_heading(doc, 'Éles telepítés (WSL2)', 2)
add_code(doc,
'# Adminként futtatva:\n'
'.\install.bat\n'
'# Vagy:\n'
'powershell -ExecutionPolicy Bypass -File "_setup\\CRT_install.ps1"'
)


# ── 9. KONFIGURÁCIÓ ───────────────────────────────────────────
add_heading(doc, '9. Konfiguráció (system_config kulcsok)', 1)
add_table(doc,
    ['Kulcs', 'Leírás', 'Hol állítható'],
    [
        ['claude_api_key',       'Anthropic API kulcs',          'beallitasok.html / .env'],
        ['claude_model',         'Alapértelmezett Claude modell', 'beallitasok.html'],
        ['smtp_host/port/…',     'Email SMTP (OTP küldéshez)',    'beallitasok.html'],
        ['company_name',         'Cég neve (dokumentumokba)',     'beallitasok.html'],
        ['ollama_url',           'Ollama szerver URL',            'beallitasok.html'],
        ['ollama_model',         'Alapértelmezett Ollama modell', 'beallitasok.html'],
        ['lora_active_job_id',   'Aktív LoRA job azonosítója',    'lora.html'],
        ['lora_adapter_path',    'LoRA adapter könyvtár',         'lora.html'],
        ['ai_conf_high/low',     'Azonosítás bizonyossági küszöb', 'beallitasok.html'],
    ],
    col_widths=[4.5, 6.5, 6]
)


# ── 10. FEJLESZTÉSI HISTÓRIA ──────────────────────────────────
add_heading(doc, '10. Fejlesztési história', 1)
add_table(doc,
    ['Verzió', 'Dátum', 'Tartalom'],
    [
        ['v0.1', '2026-05-18', 'DB séma – 15 tábla alap'],
        ['v0.2', '2026-05-18', 'FastAPI gerinc + állapotjelző'],
        ['v0.3', '2026-05-28', 'Cikktörzs API + UI + fájl parser + AI azonosítás + GitHub'],
        ['v0.4', '2026-06-08', 'Auth teljes: PIN + OTP + JWT + brute force + sanitize'],
        ['v0.5', '2026-06-08', 'Ajánlatkészítés + export + árak + naplók + beállítások'],
        ['v0.6', '2026-06-09', 'Playwright scraper + Nginx + admin init'],
        ['v0.7', '2026-06-10', 'Ollama fallback + ChromaDB embedding motor'],
        ['v0.8', '2026-06-10', 'OCR + PDF export + Golden examples + widget + Bootstrap ZIP'],
        ['v0.9', '2026-06-10', 'LLaVA Vision rajzelemző UI'],
        ['v1.0', '2026-06-10', 'LoRA finomhangolás pipeline – teljes AI lánc kész'],
    ],
    col_widths=[2, 3.5, 11.5]
)


# ── 11. LÁBJEGYZET ────────────────────────────────────────────
doc.add_paragraph()
foot_para = doc.add_paragraph()
foot_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot_para.add_run('CRT Ajánlatsegéd Whitebook v1.0  ·  Civil Rendszertechnika Kft.  ·  2026-06-10')
r.font.size = Pt(8)
r.font.color.rgb = C_DIM
r.font.name = 'Calibri'
r.italic = True


# ── MENTÉS ────────────────────────────────────────────────────
doc.save(OUT)
print(f"OK – Whitebook Word dokumentum kesz: {OUT}")
