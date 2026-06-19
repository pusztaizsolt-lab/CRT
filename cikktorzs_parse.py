"""
CRT – Cikktörzs fájl beolvasó modul
Támogatott: xlsx, xls, csv, pdf, docx, html, htm, md
"""
import io, re
from typing import List, Dict, Any

# ── NORMALIZÁLÁS ──────────────────────────────────────────────
_KEY_MAP = {
    'nev': 'name', 'név': 'name', 'megnevezés': 'name', 'megnevezes': 'name',
    'cikk': 'name', 'cikkszám': 'name', 'cikkszam': 'name', 'termék': 'name',
    'leírás': 'name', 'leiras': 'name', 'tétel': 'name', 'tetel': 'name',
    'gyártó': 'manufacturer', 'gyarto': 'manufacturer', 'brand': 'manufacturer',
    'márka': 'manufacturer', 'marka': 'manufacturer', 'szállító': 'manufacturer',
    'egység': 'unit', 'egyseg': 'unit', 'me': 'unit', 'mértékegység': 'unit',
    'kategória': 'category', 'kategoria': 'category', 'kat': 'category', 'type': 'category',
    'megjegyzés': 'notes', 'megjegyzes': 'notes', 'notes': 'notes',
}

def _normalize_row(row: Dict) -> Dict:
    result = {}
    for k, v in row.items():
        if v is None:
            continue
        s = str(v).strip()
        if not s:
            continue
        norm = _KEY_MAP.get(k.lower().strip(), None)
        if norm:
            result[norm] = s
        else:
            result[k.lower().strip()] = s
    return result

def _clean_rows(rows: List[Dict]) -> List[Dict]:
    """Üres és dupla sorok szűrése"""
    seen, out = set(), []
    for r in rows:
        name = r.get('name', '').strip()
        if len(name) < 2 or name.lower() in ('megnevezés', 'neve', 'tétel', 'leírás', 'name'):
            continue
        if name in seen:
            continue
        seen.add(name)
        out.append(r)
    return out

# ── FŐ BELÉPŐ ────────────────────────────────────────────────
def parse_file(filename: str, content: bytes) -> List[Dict[str, Any]]:
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    parsers = {
        'xlsx': parse_excel, 'xls': parse_excel,
        'csv': parse_csv,
        'pdf': parse_pdf,
        'docx': parse_word, 'doc': parse_word,
        'html': parse_html, 'htm': parse_html,
        'md': parse_markdown,
    }
    fn = parsers.get(ext)
    if not fn:
        raise ValueError(f"Nem támogatott formátum: .{ext}  (xlsx, csv, pdf, docx, html, md)")
    rows = fn(content)
    return _clean_rows(rows)


# ── EXCEL ────────────────────────────────────────────────────
def parse_excel(content: bytes) -> List[Dict]:
    try:
        import openpyxl
    except ImportError:
        raise ValueError("openpyxl csomag hiányzik – telepítés: py -3.11 -m pip install openpyxl")
    wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return []
    headers = [str(h).strip() if h is not None else f'col{i}' for i, h in enumerate(rows[0])]
    result = []
    for row in rows[1:]:
        if not any(c for c in row if c is not None):
            continue
        d = {headers[i]: row[i] for i in range(min(len(headers), len(row)))}
        n = _normalize_row({k: v for k, v in d.items() if v is not None})
        if not n.get('name') and headers:
            # Ha nincs 'name' kulcs, az első nem üres cella legyen a neve
            for h in headers:
                if d.get(h):
                    n['name'] = str(d[h]).strip()
                    break
        if n.get('name'):
            result.append(n)
    return result


# ── CSV ──────────────────────────────────────────────────────
def parse_csv(content: bytes) -> List[Dict]:
    import csv
    for enc in ('utf-8-sig', 'utf-8', 'cp1250', 'iso-8859-2'):
        try:
            text = content.decode(enc)
            break
        except:
            text = None
    if not text:
        raise ValueError("CSV kódolás nem felismerhető")
    # Elválasztó detektálás
    sample = text[:2000]
    delim = ';' if sample.count(';') > sample.count(',') else ','
    reader = csv.DictReader(io.StringIO(text), delimiter=delim)
    result = []
    for row in reader:
        n = _normalize_row(dict(row))
        if not n.get('name') and row:
            first_val = next(iter(row.values()), '')
            if first_val.strip():
                n['name'] = first_val.strip()
        if n.get('name'):
            result.append(n)
    return result


# ── PDF ──────────────────────────────────────────────────────
def parse_pdf(content: bytes) -> List[Dict]:
    try:
        import pdfplumber
    except ImportError:
        raise ValueError("pdfplumber csomag hiányzik – telepítés: py -3.11 -m pip install pdfplumber")

    result = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            # 1. Táblák elsőbbsége
            for table in (page.extract_tables() or []):
                if not table or len(table) < 2:
                    continue
                headers = [str(h).strip() if h else f'col{i}' for i, h in enumerate(table[0])]
                for row in table[1:]:
                    cells = [str(c).strip() if c else '' for c in row]
                    d = {headers[i]: cells[i] for i in range(min(len(headers), len(cells)))}
                    n = _normalize_row(d)
                    if not n.get('name') and cells:
                        n['name'] = cells[0]
                    if n.get('name'):
                        result.append(n)

            # 2. Sima szöveg (digitális PDF)
            if not result:
                txt = (page.extract_text() or '').strip()
                if txt:
                    for line in txt.split('\n'):
                        line = line.strip()
                        if len(line) > 3 and not re.match(r'^\d+$', line):
                            result.append({'name': line, '_raw': True})

            # 3. OCR fallback – képes PDF oldal (ha nincs kinyerhető szöveg)
            if not result:
                result += _ocr_page(page)

    return result


def _ocr_page(page) -> List[Dict]:
    """
    Képes PDF oldal OCR-rel (pytesseract + Pillow).
    Csak akkor hívódik, ha pdfplumber nem talált szöveget.
    """
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return []

    try:
        img = page.to_image(resolution=200).original
        if not isinstance(img, Image.Image):
            img = Image.fromarray(img)
        txt = pytesseract.image_to_string(img, lang="hun+eng", config="--psm 6")
    except Exception:
        return []

    rows = []
    for line in txt.split('\n'):
        line = line.strip()
        # Kizárjuk a rövid, szám-csak, és fejléc-szerű sorokat
        if len(line) < 4:
            continue
        if re.match(r'^[\d\s.,/-]+$', line):
            continue
        rows.append({'name': line, '_raw': True, '_ocr': True})
    return rows


# ── WORD ─────────────────────────────────────────────────────
def parse_word(content: bytes) -> List[Dict]:
    try:
        from docx import Document
    except ImportError:
        raise ValueError("python-docx csomag hiányzik – telepítés: py -3.11 -m pip install python-docx")
    doc = Document(io.BytesIO(content))
    result = []
    for table in doc.tables:
        if not table.rows:
            continue
        headers = [c.text.strip() for c in table.rows[0].cells]
        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            d = {headers[i]: cells[i] for i in range(min(len(headers), len(cells)))}
            n = _normalize_row(d)
            if not n.get('name') and cells:
                n['name'] = cells[0]
            if n.get('name'):
                result.append(n)
    # Bekezdések fallback
    if not result:
        for para in doc.paragraphs:
            t = para.text.strip()
            if len(t) > 3:
                result.append({'name': t, '_raw': True})
    return result


# ── HTML / HTM ───────────────────────────────────────────────
def parse_html(content: bytes) -> List[Dict]:
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise ValueError("beautifulsoup4 csomag hiányzik – telepítés: py -3.11 -m pip install beautifulsoup4")
    for enc in ('utf-8', 'cp1250', 'iso-8859-2'):
        try:
            text = content.decode(enc)
            break
        except:
            text = content.decode('utf-8', errors='replace')
    soup = BeautifulSoup(text, 'html.parser')
    result = []
    for table in soup.find_all('table'):
        ths = table.find_all('th')
        headers = [th.get_text(strip=True) for th in ths]
        data_rows = table.find_all('tr')
        if not headers and data_rows:
            headers = [td.get_text(strip=True) for td in data_rows[0].find_all('td')]
            data_rows = data_rows[1:]
        for tr in data_rows:
            cells = [td.get_text(strip=True) for td in tr.find_all('td')]
            if not cells:
                continue
            d = {headers[i] if i < len(headers) else f'col{i}': cells[i] for i in range(len(cells))}
            n = _normalize_row(d)
            if not n.get('name') and cells:
                n['name'] = cells[0]
            if n.get('name'):
                result.append(n)
    if not result:
        for line in soup.get_text('\n').split('\n'):
            line = line.strip()
            if len(line) > 3:
                result.append({'name': line, '_raw': True})
    return result


# ── MARKDOWN ─────────────────────────────────────────────────
def parse_markdown(content: bytes) -> List[Dict]:
    text = content.decode('utf-8', errors='replace')
    result = []
    lines = text.split('\n')
    headers: List[str] = []
    in_table = False
    for line in lines:
        stripped = line.strip()
        if '|' in stripped:
            parts = [p.strip() for p in stripped.strip('|').split('|')]
            if all(re.match(r'^[-: ]+$', p) for p in parts if p):
                in_table = True
                continue
            if not headers:
                headers = parts
                in_table = True
            elif in_table:
                d = {headers[i]: parts[i] for i in range(min(len(headers), len(parts)))}
                n = _normalize_row(d)
                if not n.get('name') and parts:
                    n['name'] = parts[0]
                if n.get('name'):
                    result.append(n)
        else:
            in_table = False
            headers = []
            m = re.match(r'^[-*+]\s+(.+)', stripped)
            if m:
                result.append({'name': m.group(1).strip(), '_raw': True})
    return result
