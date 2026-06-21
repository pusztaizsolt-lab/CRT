"""
CRT Előfeldolgozó Router — 1. réteg  [BETA]
============================================
Feladata: MIT LÁTOK ÉS MERRE MEGYEK?

  INPUT  →  [típusdetektálás]  →  [útválasztás]  →  [kinyerés]  →  PreprocessResult
             pdf_native?                pdfplumber
             pdf_scanned?      →        LLaVA OCR
             excel?            →        openpyxl
             word?             →        python-docx
             html?             →        BeautifulSoup
             kép/fotó?         →        LLaVA OCR

NEM foglalkozik cikktörzs-egyeztetéssel — az a 2. réteg (auto_match_loop) feladata.
A 2. réteg minden web_prices sort ÖNÁLLÓAN kap, nem tudja "honnan jött" az adat.
"""

import io, re, json, logging, base64, urllib.request
from dataclasses import dataclass, field
from enum import Enum
from typing import Optional

log = logging.getLogger("CRT.preprocessor")

OLLAMA_URL  = "http://localhost:11434/api/generate"
LLAVA_MODEL = "llava:7b"
LLAMA_MODEL = "llama3:8b"

# ── TÍPUSOK ──────────────────────────────────────────────────────────────────

class ContentType(str, Enum):
    EXCEL       = "excel"
    PDF_NATIVE  = "pdf_native"    # natív szöveges PDF
    PDF_SCANNED = "pdf_scanned"   # szkennelt/képes PDF → OCR kell
    WORD        = "word"
    HTML        = "html"
    IMAGE       = "image"         # JPG/PNG/WEBP fotó
    UNKNOWN     = "unknown"


class DocType(str, Enum):
    PRICE_LIST = "price_list"   # szállítói árlista — feldolgozzuk
    CATALOG    = "catalog"      # termékleírás/katalógus — részben használható
    INVOICE    = "invoice"      # számla — esetleg árinfo
    SPEC       = "spec_sheet"   # műszaki lap — általában nem tartalmaz árat
    UNKNOWN    = "unknown"      # nem osztályozható


@dataclass
class PriceItem:
    raw_name:  str
    raw_price: Optional[float] = None
    unit:      Optional[str]   = None
    sku:       Optional[str]   = None
    currency:  str             = "HUF"


@dataclass
class PreprocessResult:
    items:           list[PriceItem] = field(default_factory=list)
    content_type:    ContentType     = ContentType.UNKNOWN
    doc_type:        DocType         = DocType.UNKNOWN
    motor:           str             = "none"
    page_count:      int             = 0
    text_chars:      int             = 0
    warnings:        list[str]       = field(default_factory=list)
    raw_text_sample: str             = ""   # első 500 char — debug/audit


# ── TÍPUSDETEKTÁLÁS ───────────────────────────────────────────────────────────

_EXCEL_MAGIC = (b"PK\x03\x04", b"\xd0\xcf\x11\xe0")  # OOXML + OLE2
_PDF_MAGIC   = b"%PDF"
_DOCX_MAGIC  = b"PK\x03\x04"  # ugyanaz mint xlsx — el kell különíteni filename alapján

def detect_content_type(file_bytes: bytes, filename: str) -> ContentType:
    """
    Tartalomtípus detektálás — kétlépéses:
      1. fájl magic bytes (nem csak kiterjesztés)
      2. kiterjesztés segít az ambiguitásnál (OOXML xlsx vs docx)
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    hdr = file_bytes[:8]

    if hdr[:4] == _PDF_MAGIC:
        is_scanned = _is_pdf_scanned(file_bytes)
        return ContentType.PDF_SCANNED if is_scanned else ContentType.PDF_NATIVE

    if hdr[:4] == b"PK\x03\x04":
        if ext in ("xlsx", "xls", "xlsm", "xlsb"):
            return ContentType.EXCEL
        if ext in ("docx", "doc"):
            return ContentType.WORD
        # ambiguous — kiterjesztés alapján döntünk
        return ContentType.EXCEL if ext.startswith("xl") else ContentType.WORD

    if hdr[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":  # OLE2 (régi xls/doc)
        return ContentType.EXCEL if ext in ("xls", "xlsm") else ContentType.WORD

    if file_bytes[:6] in (b"\xff\xd8\xff\xe0", b"\xff\xd8\xff\xe1"):
        return ContentType.IMAGE  # JPEG
    if file_bytes[:8] == b"\x89PNG\r\n\x1a\n":
        return ContentType.IMAGE  # PNG
    if file_bytes[:4] in (b"RIFF", b"WEBP"):
        return ContentType.IMAGE  # WEBP

    if b"<html" in file_bytes[:2000].lower() or b"<!doctype" in file_bytes[:200].lower():
        return ContentType.HTML

    return ContentType.UNKNOWN


def _is_pdf_scanned(pdf_bytes: bytes) -> bool:
    """
    Ha pdfplumber nem tud szöveget kinyerni → szkennelt.
    Legalább 50 karakter kell az első 5 oldalon.
    """
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            chars = 0
            for page in pdf.pages[:5]:
                t = page.extract_text() or ""
                chars += len(t.strip())
                if chars >= 50:
                    return False
        return True
    except Exception as e:
        log.warning("PDF szkennelt-detektálás hiba: %s", e)
        return False  # biztonságos alap: native útvonalon próbálja


# ── DOKUMENTUM OSZTÁLYOZÓ ─────────────────────────────────────────────────────

_PRICE_KEYWORDS = re.compile(
    r"\b(árlista|price\s*list|nettó|bruttó|Ft\b|HUF\b|EUR\b|egységár|nett\s*price"
    r"|szállítói|listaár|mennyiség|katalógusár|article|cikkszám|sku\b)",
    re.I
)

def classify_document(text_sample: str) -> DocType:
    """
    Gyors osztályozás regex alapján — llama3 inferencia nélkül.
    Ha a szöveg ≥3 ár-jellegű kulcsszót tartalmaz → price_list.
    Fallback: llama3 osztályozás (ha Ollama fut).
    """
    hits = len(_PRICE_KEYWORDS.findall(text_sample))
    if hits >= 3:
        return DocType.PRICE_LIST

    # Llama3 osztályozás ha regex nem döntött
    try:
        prompt = (
            "Kategorizáld ezt a dokumentumot EGY szóval:\n"
            "price_list | catalog | invoice | spec_sheet | unknown\n\n"
            f"Szöveg (részlet):\n{text_sample[:600]}\n\nVálasz:"
        )
        resp = _ollama_generate(prompt, model=LLAMA_MODEL, max_tokens=10)
        resp = resp.lower().strip()
        if "price" in resp or "árlista" in resp:
            return DocType.PRICE_LIST
        if "catalog" in resp:
            return DocType.CATALOG
        if "invoice" in resp or "számla" in resp:
            return DocType.INVOICE
        if "spec" in resp:
            return DocType.SPEC
    except Exception:
        pass

    return DocType.UNKNOWN


# ── SZÖVEG → STRUKTURÁLT LISTA ────────────────────────────────────────────────

_PRICE_LINE = re.compile(
    r"(?P<name>[^\|]{4,120}?)"
    r"\s*[\|;]\s*"
    r"(?P<price>[\d\s.,]+)"
    r"\s*(?P<currency>Ft|HUF|EUR|USD|€)?"
    r"\s*(?:nett[oó]|bruttó|net)?",
    re.I
)

def _parse_price_text(text: str) -> list[PriceItem]:
    """
    Szövegből strukturált ársorok kinyerése.
    Elsőként regex (gyors), utána llama3 ha kevés találat.
    """
    items = _regex_parse(text)
    if len(items) >= 3:
        return items
    # Regex keveset talált — llama3 próba (max 30s)
    try:
        chunk = text[:5000]
        raw   = _ollama_generate(
            "Magyar szállítói árlista részlete.\n"
            "Listazd az összes terméket és nettó árat.\n"
            "Formátum soronként: TERMÉKNÉV | ÁR Ft\n"
            "Ha nincs ár a tételnél, hagyd ki. Ha nincs semmi, írj: NINCS\n\n"
            + chunk,
            max_tokens=1000,
            timeout=30,
        )
        items = _regex_parse(raw)
    except Exception as e:
        log.debug("llama3 kinyerés hiba: %s", e)
    return items


def _regex_parse(text: str) -> list[PriceItem]:
    items = []
    for line in text.splitlines():
        line = line.strip()
        if len(line) < 6:
            continue
        m = _PRICE_LINE.search(line)
        if m:
            name  = m.group("name").strip(" \t-–•*")
            price = _parse_price(m.group("price"))
            cur   = (m.group("currency") or "HUF").replace("€","EUR").upper()
            if name and price and price > 0:
                items.append(PriceItem(raw_name=name[:200], raw_price=price, currency=cur))
            continue
        # Ft/EUR szám nélkül | formátumban — egyszerűbb parse
        if re.search(r'\d[\d\s,.]*\s*(Ft|HUF|EUR|€)', line, re.I):
            nums = re.findall(r'[\d\s,.]+', line)
            for n in nums:
                price = _parse_price(n)
                if price and price > 0:
                    name_part = re.sub(r'[\d\s,./()%Ft€HUEUR]+', ' ', line).strip(" \t-–•")
                    if len(name_part) > 3:
                        items.append(PriceItem(raw_name=name_part[:200], raw_price=price))
                    break
    return items


def _parse_price(s: str) -> Optional[float]:
    try:
        s = s.replace("\xa0", "").replace(" ", "").strip()
        # magyar formátum: 1.234,56 → 1234.56
        if re.match(r"^\d{1,3}(\.\d{3})+(,\d{1,4})?$", s):
            s = s.replace(".", "").replace(",", ".")
        elif re.match(r"^\d{1,3}(\s\d{3})+(,\d{1,4})?$", s):
            s = s.replace(" ", "").replace(",", ".")
        else:
            s = s.replace(",", ".")
        v = float(s)
        return v if 0.01 <= v <= 99_999_999 else None
    except (ValueError, AttributeError):
        return None


# ── OLLAMA HÍVÓK ─────────────────────────────────────────────────────────────

def _ollama_generate(prompt: str, model: str = LLAMA_MODEL,
                     max_tokens: int = 800, timeout: int = 120) -> str:
    payload = json.dumps({
        "model":   model,
        "prompt":  prompt,
        "stream":  False,
        "options": {"temperature": 0.0, "num_predict": max_tokens},
    }).encode()
    req = urllib.request.Request(
        OLLAMA_URL, data=payload,
        headers={"Content-Type": "application/json"}
    )
    with urllib.request.urlopen(req, timeout=timeout) as r:
        return json.loads(r.read())["response"]


def _ask_llama_prices(text: str) -> str:
    return _ollama_generate(
        "Magyar szállítói árlista részlete.\n"
        "Listazd az összes terméket és nettó árat.\n"
        "Formátum soronként: TERMÉKNÉV | ÁR Ft\n"
        "Ha nincs ár a tételnél, hagyd ki. Ha nincs semmi, írj: NINCS\n\n"
        + text,
        max_tokens=1000,
    )


def _llava_ocr(image_bytes: bytes, page_hint: str = "") -> str:
    """LLaVA multimodal OCR — kép vagy PDF-oldal → szöveg."""
    b64 = base64.b64encode(image_bytes).decode()
    prompt = (
        f"Magyar szállítói árlista{' – ' + page_hint if page_hint else ''}.\n"
        "Olvasd ki PONTOSAN az összes terméknevet és árat.\n"
        "Formátum soronként: TERMÉKNÉV | ÁR Ft\n"
        "Ha nem látsz árat, csak a nevet írd ki egyszerű sorként.\n"
        "Ha az oldal nem tartalmaz termékeket: ÜRES"
    )
    payload = json.dumps({
        "model":   LLAVA_MODEL,
        "prompt":  prompt,
        "images":  [b64],
        "stream":  False,
        "options": {"temperature": 0.0, "num_predict": 1200},
    }).encode()
    req = urllib.request.Request(
        OLLAMA_URL, data=payload,
        headers={"Content-Type": "application/json"}
    )
    with urllib.request.urlopen(req, timeout=180) as r:
        return json.loads(r.read())["response"]


def _pdf_to_images(pdf_bytes: bytes) -> list[bytes]:
    """
    PDF oldalak → PNG képek lista.
    Próbálja: pymupdf (fitz) → pdf2image → None.
    """
    # 1. pymupdf (fitz) — gyors, C extension
    try:
        import fitz  # pymupdf
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        images = []
        for page in doc:
            mat = fitz.Matrix(2.0, 2.0)  # 2x zoom → jobb OCR minőség
            pix = page.get_pixmap(matrix=mat)
            images.append(pix.tobytes("png"))
        doc.close()
        return images
    except ImportError:
        pass
    except Exception as e:
        log.warning("pymupdf hiba: %s", e)

    # 2. pdf2image — poppler szükséges
    try:
        from pdf2image import convert_from_bytes
        pages = convert_from_bytes(pdf_bytes, dpi=200)
        images = []
        for page in pages:
            buf = io.BytesIO()
            page.save(buf, format="PNG")
            images.append(buf.getvalue())
        return images
    except ImportError:
        pass
    except Exception as e:
        log.warning("pdf2image hiba: %s", e)

    return []  # nem sikerült — upstream kezel


# ── EXTRAKCIÓS ÚTVONALAK ──────────────────────────────────────────────────────

def _extract_excel(file_bytes: bytes) -> tuple[list[PriceItem], str, int]:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    items: list[PriceItem] = []
    debug_lines: list[str] = []
    _NAME_HINTS  = re.compile(r"nev|name|megnevez|termek|leiras|desc", re.I)
    _PRICE_HINTS = re.compile(r"ar|price|netto|brutto|nett[oó]|ft|huf|eur", re.I)
    _CODE_HINTS  = re.compile(r"cikk|kod|code|szam|sku|id|ref", re.I)

    for sheet in wb.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue
        # Fejlécsor keresése (első nem-üres sor)
        header = None
        header_idx = 0
        for i, row in enumerate(rows[:5]):
            cells = [str(c or "").strip() for c in row]
            if any(c for c in cells):
                header = cells
                header_idx = i
                break
        if header is None:
            continue

        # Oszlop-index meghatározás fejléc alapján
        name_col  = next((i for i, h in enumerate(header) if _NAME_HINTS.search(h)), None)
        price_col = next((i for i, h in enumerate(header) if _PRICE_HINTS.search(h)), None)
        code_col  = next((i for i, h in enumerate(header) if _CODE_HINTS.search(h)), None)

        # Fallback: ha nincs fejléc-találat, próbáljuk pozíció alapján (0=kód,1=név,3=ár)
        if name_col is None and len(header) >= 2:
            name_col = 1
        if price_col is None and len(header) >= 4:
            price_col = 3

        for row in rows[header_idx + 1:]:
            cells = [str(c or "").strip() for c in row]
            if not any(c for c in cells):
                continue
            name  = cells[name_col]  if name_col  is not None and name_col  < len(cells) else ""
            price_raw = cells[price_col] if price_col is not None and price_col < len(cells) else ""
            code  = cells[code_col]  if code_col  is not None and code_col  < len(cells) else ""
            price = _parse_price(price_raw)
            if name and price and price > 0:
                items.append(PriceItem(
                    raw_name=f"{code} {name}".strip()[:200] if code else name[:200],
                    raw_price=price,
                ))
                debug_lines.append(f"{name} | {price} Ft")
            else:
                debug_lines.append("  |  ".join(c for c in cells if c))
        if len(debug_lines) > 3000:
            break

    wb.close()
    # Ha a fejlécalapú kinyerés sikertelen, fallback a regex/llama útra
    if not items:
        text = "\n".join(debug_lines[:3000])
        items = _parse_price_text(text)
    else:
        text = "\n".join(debug_lines[:3000])
    return items, text, 1


def _extract_pdf_native(file_bytes: bytes) -> tuple[list[PriceItem], str, int]:
    import pdfplumber
    all_text = []
    page_cnt = 0
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        page_cnt = len(pdf.pages)
        for page in pdf.pages[:30]:
            # Táblázat kinyerés elsőként (strukturáltabb)
            tables = page.extract_tables() or []
            for tbl in tables:
                for row in tbl:
                    cells = [str(c or "").strip() for c in row if c]
                    if cells:
                        all_text.append("  |  ".join(cells))
            # Szöveg ha nincs táblázat
            if not tables:
                t = page.extract_text() or ""
                all_text.extend(t.splitlines())
            if len(all_text) > 3000:
                break

    text  = "\n".join(all_text[:3000])
    items = _parse_price_text(text)
    return items, text, page_cnt


def _extract_pdf_scanned(file_bytes: bytes, warnings: list[str]) -> tuple[list[PriceItem], str, int]:
    images = _pdf_to_images(file_bytes)
    if not images:
        warnings.append(
            "PDF→kép konverzió sikertelen (pip install pymupdf VAGY pdf2image + poppler). "
            "Próbáld natív PDF úton."
        )
        # Fallback: native path (esetleg mégis van valami szöveg)
        return _extract_pdf_native(file_bytes)

    all_text = []
    for i, img in enumerate(images[:15]):  # max 15 oldal OCR-ban
        try:
            ocr_text = _llava_ocr(img, page_hint=f"{i+1}. oldal")
            if "ÜRES" not in ocr_text.upper():
                all_text.append(ocr_text)
            log.info("LLaVA OCR oldal %d/%d: %d char", i+1, len(images), len(ocr_text))
        except Exception as e:
            warnings.append(f"LLaVA OCR {i+1}. oldal hiba: {e}")

    text  = "\n".join(all_text)
    items = _parse_price_text(text)
    return items, text, len(images)


def _extract_word(file_bytes: bytes) -> tuple[list[PriceItem], str, int]:
    import docx
    doc   = docx.Document(io.BytesIO(file_bytes))
    lines = []

    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append("  |  ".join(cells))

    text  = "\n".join(lines[:3000])
    items = _parse_price_text(text)
    return items, text, 1


def _extract_image(file_bytes: bytes, warnings: list[str]) -> tuple[list[PriceItem], str, int]:
    try:
        ocr_text = _llava_ocr(file_bytes)
        items    = _parse_price_text(ocr_text)
        return items, ocr_text, 1
    except Exception as e:
        warnings.append(f"LLaVA OCR hiba: {e} — ellenőrizd hogy fut-e a llava:7b modell")
        return [], "", 0


def _extract_html(file_bytes: bytes) -> tuple[list[PriceItem], str, int]:
    html = file_bytes.decode("utf-8", errors="ignore")
    # Script és style eltávolítás
    html = re.sub(r'<script[^>]*>.*?</script>', ' ', html, flags=re.S | re.I)
    html = re.sub(r'<style[^>]*>.*?</style>',  ' ', html, flags=re.S | re.I)
    html = re.sub(r'<[^>]+>', ' ', html)
    html = re.sub(r'&nbsp;', ' ', html)
    html = re.sub(r'&amp;',  '&', html)
    html = re.sub(r'\s+', ' ', html).strip()
    items = _parse_price_text(html)
    return items, html[:5000], 1


# ── FŐ BELÉPŐPONT ─────────────────────────────────────────────────────────────

def process(file_bytes: bytes, filename: str) -> PreprocessResult:
    """
    Fő belépőpont — automatikusan felismeri a típust és kinyeri a tételeket.

    Returns PreprocessResult with items, content_type, doc_type, motor, warnings.
    """
    result   = PreprocessResult()
    warnings = result.warnings

    # 1. Típusdetektálás
    ct = detect_content_type(file_bytes, filename)
    result.content_type = ct
    log.info("Preprocessor: %s → %s", filename, ct.value)

    # 2. Kinyerés az útvonal alapján
    try:
        if ct == ContentType.EXCEL:
            items, text, pages = _extract_excel(file_bytes)
            result.motor = "openpyxl+llama3"

        elif ct == ContentType.PDF_NATIVE:
            items, text, pages = _extract_pdf_native(file_bytes)
            result.motor = "pdfplumber+llama3"

        elif ct == ContentType.PDF_SCANNED:
            log.info("Szkennelt PDF — LLaVA OCR indul (%s)", filename)
            items, text, pages = _extract_pdf_scanned(file_bytes, warnings)
            result.motor = "llava_ocr"

        elif ct == ContentType.WORD:
            items, text, pages = _extract_word(file_bytes)
            result.motor = "python-docx+llama3"

        elif ct == ContentType.IMAGE:
            log.info("Kép OCR — LLaVA indul (%s)", filename)
            items, text, pages = _extract_image(file_bytes, warnings)
            result.motor = "llava_ocr"

        elif ct == ContentType.HTML:
            items, text, pages = _extract_html(file_bytes)
            result.motor = "html_parser+llama3"

        else:
            warnings.append(f"Ismeretlen fájltípus: {filename} — feldolgozás kihagyva")
            return result

        result.items        = items
        result.page_count   = pages
        result.text_chars   = len(text)
        result.raw_text_sample = text[:500]

    except ImportError as e:
        warnings.append(
            f"Hiányzó könyvtár: {e} — "
            f"pip install pdfplumber openpyxl python-docx pymupdf"
        )
        return result
    except Exception as e:
        warnings.append(f"Kinyerési hiba ({ct.value}): {e}")
        log.error("Preprocessor kinyerési hiba: %s", e, exc_info=True)
        return result

    # 3. Dokumentum-osztályozás
    result.doc_type = classify_document(result.raw_text_sample)

    log.info(
        "Preprocessor kész: %s | %s | motor=%s | %d tétel | %d char | %d oldal",
        ct.value, result.doc_type.value, result.motor,
        len(items), result.text_chars, result.page_count,
    )
    if warnings:
        log.warning("Preprocessor figyelmeztetések: %s", "; ".join(warnings))

    return result


# ── STANDALONE TESZT ─────────────────────────────────────────────────────────

if __name__ == "__main__":
    import sys, pathlib
    logging.basicConfig(level=logging.INFO,
                        format="%(asctime)s | %(levelname)s | %(message)s")

    if len(sys.argv) < 2:
        print("Használat: py -3.11 preprocessor.py <fájlútvonal>")
        sys.exit(1)

    path  = pathlib.Path(sys.argv[1])
    data  = path.read_bytes()
    res   = process(data, path.name)

    print(f"\n{'='*60}")
    print(f"Fájl:       {path.name}")
    print(f"Típus:      {res.content_type.value}")
    print(f"Dokumentum: {res.doc_type.value}")
    print(f"Motor:      {res.motor}")
    print(f"Oldalak:    {res.page_count}")
    print(f"Szöveg:     {res.text_chars} char")
    print(f"Tételek:    {len(res.items)}")
    if res.warnings:
        for w in res.warnings:
            print(f"[!] {w}")
    print(f"{'='*60}")
    for i, it in enumerate(res.items[:20]):
        print(f"  {i+1:3d}. {it.raw_name[:55]:55s} | {it.raw_price or '–':>12} {it.currency}")
    if len(res.items) > 20:
        print(f"  … és még {len(res.items)-20} tétel")
    print(f"\nSzövegminta:\n{res.raw_text_sample[:300]}")
